"""Document parsing — uses LlamaParse for PDF/DOCX, native readers for plain text."""
from __future__ import annotations

import logging
from pathlib import Path

from llama_index.core import Document

from .config import get_settings
from .usage import get_counter

logger = logging.getLogger(__name__)

PLAIN_TEXT_EXT = {".txt", ".md", ".csv", ".json", ".log"}
LLAMAPARSE_EXT = {".pdf", ".docx", ".doc", ".pptx", ".ppt", ".xlsx", ".xls", ".html", ".htm"}


def _read_plain(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [Document(text=text, metadata={"file_name": path.name})]


def _read_with_llamaparse(path: Path) -> list[Document]:
    """Use LlamaParse for rich documents. Falls back to pypdf / python-docx on failure."""
    settings = get_settings()
    try:
        from llama_parse import LlamaParse  # noqa: WPS433

        parser = LlamaParse(
            api_key=settings.llama_cloud_api_key,
            result_type="markdown",
            verbose=False,
            language="en",
        )
        docs = parser.load_data(str(path))
        for d in docs:
            d.metadata = {**(d.metadata or {}), "file_name": path.name}
        if docs and any(d.text and d.text.strip() for d in docs):
            get_counter().increment("llamaparse_pages", len(docs))
            return docs
        logger.warning("LlamaParse returned empty docs for %s, falling back", path.name)
    except Exception as exc:  # noqa: BLE001
        logger.warning("LlamaParse failed for %s: %s — falling back", path.name, exc)

    return _fallback_parse(path)


def _fallback_parse(path: Path) -> list[Document]:
    ext = path.suffix.lower()
    text = ""
    try:
        if ext == ".pdf":
            from pypdf import PdfReader  # noqa: WPS433

            reader = PdfReader(str(path))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        elif ext in {".docx"}:
            from docx import Document as DocxDocument  # noqa: WPS433

            doc = DocxDocument(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        logger.error("Fallback parser failed for %s: %s", path.name, exc)
        text = ""

    return [Document(text=text or f"[Empty document: {path.name}]", metadata={"file_name": path.name})]


def parse_file(path: Path) -> list[Document]:
    """Parse a file into LlamaIndex Documents."""
    ext = path.suffix.lower()
    if ext in PLAIN_TEXT_EXT:
        return _read_plain(path)
    if ext in LLAMAPARSE_EXT:
        return _read_with_llamaparse(path)
    # Unknown extension — try plain text
    return _read_plain(path)
